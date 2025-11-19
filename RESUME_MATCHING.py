#!/usr/bin/env python
# coding: utf-8

# In[1]:


import os
print("Python:", __import__('sys').executable)
print("OPENROUTER_API_KEY:", os.getenv("OPENROUTER_API_KEY"))


# Importing all required libraries

# In[1]:


import os, io, re, json, time, hashlib
from collections import defaultdict
import numpy as np
import pandas as pd
import requests
from IPython.display import display, clear_output, HTML
import ipywidgets as widgets

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import cm
    import base64
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    import pdfplumber
except:
    pdfplumber = None
try:
    import docx
except:
    docx = None

print("Libraries imported.")
if not REPORTLAB_AVAILABLE:
    print(" Reportlab is  not installed.")


# SETTING API KEYS AND MODELS

# In[2]:


OPENROUTER_KEY = os.getenv("OPENROUTER_API_KEY") #api key is set in enviornment
OPENROUTER_EMBED_URL = "https://openrouter.ai/api/v1/embeddings"
OPENROUTER_CHAT_URL = "https://openrouter.ai/api/v1/chat/completions"

EMBED_CANDIDATES = ["openai/text-embedding-3-large","openai/text-embedding-3-small"]    #this model is used for embedding
LLM_EXTRACT_MODEL = "openai/gpt-3.5-turbo"   #extract structured clusters
CHAT_MODEL_DEFAULT = "mistralai/mistral-7b-instruct"  #used for generating explanations

SEMANTIC_THRESHOLD = 0.75      #setting threshold to consider a skill as matched
KEYWORD_BONUS_THRESHOLD = 0.7  
print("Configuration set.")


# FILE PARSING,CLEANING

# In[3]:


def extract_text_from_bytes(filename, b):  #extract contents if uploaded file is pdf, txt or docs
    fname = (filename or "").lower()
    if fname.endswith(".txt"):
        return b.decode("utf-8", errors="replace")
    if fname.endswith(".pdf") and pdfplumber:
        try:
            pages=[]
            with pdfplumber.open(io.BytesIO(b)) as pdf:
                for p in pdf.pages:
                    pages.append(p.extract_text() or "")
            return "\n".join(pages)
        except:
            return ""
  
    if fname.endswith(".docx") and docx:
        try:
            doc = docx.Document(io.BytesIO(b))
            return "\n".join([p.text for p in doc.paragraphs])
        except:
            return ""
    if fname.endswith(".csv"):
        try:
            df = pd.read_csv(io.BytesIO(b))
            text_cols = [c for c in df.columns if 'text' in c.lower() or 'resume' in c.lower() or 'cv' in c.lower()]
            if text_cols:
                return "\n\n".join(df[text_cols[0]].astype(str).tolist())
            return df.to_csv(index=False)
        except:
            return b.decode("utf-8", errors="replace")
    try:
        return b.decode("utf-8", errors="replace")
    except:
        return ""

def normalize_fileupload_value(val):
    out=[]
    if not val:
        return out
    if isinstance(val, dict):
        for k,v in val.items():
            if isinstance(v, dict) and 'content' in v:
                name = v.get('name') or k
                content = v.get('content')
                if isinstance(content, (memoryview, bytearray)):
                    content = bytes(content)
                out.append({"name":name,"content":content})
    elif isinstance(val, (list,tuple)):
        for item in val:
            if isinstance(item, dict) and 'content' in item:
                name = item.get('name') or "uploaded"
                content = item.get('content')
                if isinstance(content, (memoryview, bytearray)):
                    content = bytes(content)
                out.append({"name":name,"content":content})
    return out

def sent_tokenize(text):
    text = (text or "").strip()
    if not text:
        return []
    sents = re.split(r'(?<=[.!?])\s+', text)
    return [s.strip() for s in sents if s.strip()]

def clean_text(txt):
    txt = str(txt or "")
    txt = txt.replace("\x0c"," ")
    txt = re.sub(r'\(cid:\d+\)',' ', txt)
    txt = re.sub(r'\s+',' ', txt).strip()
    return txt
    
print("File parsing and text cleaning functions done.")


# SETTING OPENROUTER API COMMUNICATION

# In[4]:


def detect_working_openrouter_embedding_model(candidates=EMBED_CANDIDATES, test_text="hello"):
    if not OPENROUTER_KEY:
        print("OPENROUTER_API_KEY not set in environment.")
        return None
    headers = {"Authorization": f"Bearer {OPENROUTER_KEY}", "Content-Type":"application/json"}
    for m in candidates:
        payload = {"model": m, "input": [test_text]}
        try:
            resp = requests.post(OPENROUTER_EMBED_URL, headers=headers, json=payload, timeout=20)
            if resp.status_code == 200:
                j = resp.json()
                if "data" in j and isinstance(j["data"], list) and "embedding" in j["data"][0]:
                    return m
        except Exception:
            continue
    return None

def get_openrouter_embeddings(texts, model, batch_size=16):
    if not OPENROUTER_KEY:
        raise RuntimeError("OPENROUTER_API_KEY missing.")
    headers = {"Authorization": f"Bearer {OPENROUTER_KEY}", "Content-Type":"application/json"}
    all_embs = []
    for i in range(0, len(texts), batch_size):
        batch = texts[i:i+batch_size]
        payload = {"model": model, "input": batch}
        try:
            resp = requests.post(OPENROUTER_EMBED_URL, headers=headers, json=payload, timeout=60)
            resp.raise_for_status()
        except requests.exceptions.RequestException as e:
            try:
                print("OpenRouter error:", resp.status_code, resp.json())
            except:
                pass
            raise RuntimeError(f"OpenRouter embeddings failed: {e}")
            
        j = resp.json()
        for item in j["data"]:
            all_embs.append(np.array(item["embedding"], dtype=np.float32))
        time.sleep(0.05)
    if not all_embs:
        return np.zeros((0,0), dtype=np.float32)
    arr = np.vstack(all_embs)
    arr = arr / np.linalg.norm(arr, axis=1, keepdims=True)
    return arr

def openrouter_chat(prompt, model=CHAT_MODEL_DEFAULT, temperature=0.0, max_tokens=512):
    if not OPENROUTER_KEY:
        raise RuntimeError("OPENROUTER_API_KEY missing.")
    headers = {"Authorization": f"Bearer {OPENROUTER_KEY}", "Content-Type":"application/json"}
    payload = {"model": model, "messages": [{"role":"user","content": prompt}], "temperature": temperature, "max_tokens": max_tokens}
    resp = requests.post(OPENROUTER_CHAT_URL, headers=headers, json=payload, timeout=120)
    resp.raise_for_status()
    j = resp.json()
    return j["choices"][0]["message"]["content"].strip()
    
print("OpenRouter API functions defined.")


# EXTRACTIONS AND MATCHING

# In[5]:


def extract_requirement_clusters(job_text, model=LLM_EXTRACT_MODEL):   #specified that o/p should be in json in the prompt
    prompt = (
        "You are a concise parser. Given a job description, return a JSON object with keys "
        "\"must_have\", \"important\", \"nice_to_have\". Each value should be an array of short skill/competency phrases (1-5 words each). "
        "It is CRITICAL that you extract the actual skills/competencies, not filler words like 'of' or 'the'. "
        "If the job contains explicit labels like 'must have' or lists, follow them. Keep phrases short. Return only JSON.\n\n"
        f"JOB_DESCRIPTION:\n{job_text}\n\nReturn JSON now."
    )
    txt = openrouter_chat(prompt, model=model, temperature=0.0, max_tokens=500)
    m = re.search(r"\{.*\}", txt, flags=re.S)
    
    if not m:
        print("Warning: LLM failed to return structured JSON. Using token frequency fallback.")
        words = re.findall(r"\b[A-Za-z0-9\-\+#\.\_]+\b", job_text.lower())
        common = sorted(set(words), key=lambda w: -words.count(w))[:12]
        return {"must_have": common[:4], "important": common[4:8], "nice_to_have": common[8:12]}
        
    try:
        obj = json.loads(m.group(0))
    except json.JSONDecodeError:
        print("Warning: LLM returned invalid JSON. Using token frequency fallback.")
        words = re.findall(r"\b[A-Za-z0-9\-\+#\.\_]+\b", job_text.lower())
        common = sorted(set(words), key=lambda w: -words.count(w))[:12]
        obj = {"must_have": common[:4], "important": common[4:8], "nice_to_have": common[8:12]}
        
    for k in ("must_have","important","nice_to_have"):
        if k not in obj:
            obj[k] = []
        if isinstance(obj[k], list):
            obj[k] = [p for p in obj[k] if len(p.split()) > 1 or p.lower() not in ['of', 'and', 'in', 'to', 'the', 'with', 'for', 'or']]
            
    return obj

def compute_skill_matches_for_resume(resume_text, skill_phrases, embed_model, chunk_size=3):
    out = {}
    r_low = resume_text.lower()
  
    for p in skill_phrases:
        p_low = p.lower()
        kw_match = bool(re.search(r"\b" + re.escape(p_low) + r"\b", r_low))
        out[p] = {"kw_match": kw_match, "sim": 0.0, "best_evidence": None}
    
    #chunking
    sentences = sent_tokenize(resume_text)
    chunks = [" ".join(sentences[i:i+chunk_size]) for i in range(len(sentences)) if " ".join(sentences[i:i+chunk_size])]
    
    if not chunks or not skill_phrases:
        return out
        
    try:
        texts = skill_phrases + chunks
        embs = get_openrouter_embeddings(texts, model=embed_model, batch_size=8)
        
        skill_embs = embs[:len(skill_phrases)]
        chunk_embs = embs[len(skill_phrases):]
        
        sims_matrix = (skill_embs @ chunk_embs.T).astype(float) 
        
        for i, p in enumerate(skill_phrases):
            max_sim = np.max(sims_matrix[i])
            best_chunk_index = np.argmax(sims_matrix[i])
            
            out[p]["sim"] = float(max_sim)
            out[p]["best_evidence"] = {
                "skill": p,
                "chunk_text": chunks[best_chunk_index],
                "score": float(max_sim)
            }
            
    except Exception as e:
        print(f"Error during embedding computation: {e}")
        pass
        
    return out
    
print("LLM extraction and Embedding matching done")


# EXPLANATION AND RANKING

# In[6]:


def generate_candidate_comparator_explanation(candidateA, candidateB):
    scoreA = candidateA.get('overall_score') if candidateA.get('overall_score') is not None else 0.0
    scoreB = candidateB.get('overall_score') if candidateB.get('overall_score') is not None else 0.0

    explanation = (
        f"CANDIDATE A (id:{candidateA.get('id')} name:{candidateA.get('name')} score:{scoreA:.3f}):\n"
        f"Matched must-have skills: {', '.join(candidateA.get('matched_skills',[]))}\n"
        f"Gaps: {', '.join(candidateA.get('gaps',[]))}\n"
        f"Top evidence snippets:\n"
    )
    for e in candidateA.get('top_evidence', []):
        explanation += f" - **{e.get('skill', 'Skill')}:** (Score {e.get('score', 0.0):.3f}) *{e['chunk_text'][:100]}...*\n"

    explanation += f"\nCANDIDATE B (id:{candidateB.get('id')} name:{candidateB.get('name')} score:{scoreB:.3f}):\n"
    explanation += f"Matched must-have skills: {', '.join(candidateB.get('matched_skills',[]))}\n"
    explanation += f"Gaps: {', '.join(candidateB.get('gaps',[]))}\n"
    explanation += f"Top evidence snippets:\n"
    for e in candidateB.get('top_evidence', []):
        explanation += f" - **{e.get('skill', 'Skill')}:** (Score {e.get('score', 0.0):.3f}) *{e['chunk_text'][:100]}...*\n"

    if scoreA > scoreB:
        explanation += "\nConclusion: Candidate A is stronger based on overall score.\n"
    elif scoreB > scoreA:
        explanation += "\nConclusion: Candidate B is stronger based on overall score.\n"
    else:
        explanation += "\nConclusion: Both candidates have equal scores.\n"

    return explanation

def explainable_ranking_pipeline(candidates):
    for c in candidates:
        if c.get('overall_score') is None:
            c['overall_score'] = 0.0
        if 'summary' not in c:
            c['summary'] = "No summary available."
    sorted_candidates = sorted(candidates, key=lambda x: x['overall_score'], reverse=True)
    explanations = []
    
    if len(sorted_candidates) >= 2:
        for i in range(len(sorted_candidates) - 1):
            a = sorted_candidates[i]
            b = sorted_candidates[i + 1]
            explanation = generate_candidate_comparator_explanation(a, b)
            explanations.append(explanation)
            
    return explanations
    
print("Explanation and Ranking done.")


# INTEGRATING PDF OPTION

# In[7]:


def generate_pdf(job_text, requirements, candidates, explanations, filename="Candidate_Ranking.pdf"):
    if not REPORTLAB_AVAILABLE:
        print("Reportlab not installed. Skipping PDF generation.")
        return
        
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    margin = 2*cm
    y = height - margin

    def write_line(text, line_height=14, is_bold=False):
        nonlocal y
        if is_bold:
             c.setFont("Helvetica-Bold", line_height-4)
        else:
             c.setFont("Helvetica", line_height-4)

        lines = text.split('\n')
        for line in lines:
            max_width = width - 2 * margin
            current_line = ""
            for word in line.split(' '):
                test_line = current_line + word + " "
                if c.stringWidth(test_line) < max_width:
                    current_line = test_line
                else:
                    c.drawString(margin, y, current_line.strip())
                    y -= line_height
                    current_line = word + " "
                
            if current_line.strip():
                c.drawString(margin, y, current_line.strip())
                y -= line_height
                
            if y < margin:
                c.showPage()
                y = height - margin
        c.setFont("Helvetica", 12)

    c.setFont("Helvetica-Bold", 16)
    write_line("Candidate Matching Report", line_height=20, is_bold=True)
    c.setFont("Helvetica", 12)
    write_line("")
    write_line("Job Description:", 16, is_bold=True)
    write_line("")
    write_line("Extracted Requirement Clusters:", 16, is_bold=True)
    for k,v in requirements.items():
        write_line(f"{k}: {', '.join(v)}")

    write_line("")
    write_line("Candidate Comparisons:", 16, is_bold=True)
    for exp in explanations:
        write_line(exp)
        write_line("-"*50)

    write_line("")
    write_line("Ranking Summary:", 16, is_bold=True)
    for cnd in sorted(candidates, key=lambda x: x['overall_score'], reverse=True):
        write_line(f"RANK {cnd['name']} - Score: {cnd['overall_score']:.3f}")
        write_line(f"Matched Skills: {', '.join(cnd['matched_skills'])}")
        write_line(f"Gaps: {', '.join(cnd['gaps'])}")
        write_line("-" * 20)


    c.save()
    buffer.seek(0)
    b64 = base64.b64encode(buffer.read()).decode()
    href = f'<a download="{filename}" href="data:application/pdf;base64,{b64}">Download PDF Report</a>'
    display(HTML(href))

print("PDF generation function done")


# In[8]:


def on_process_click(b):
    with output:
        clear_output()
        job_text = job_textarea.value.strip()
        if not job_text and job_upload.value:
            files = normalize_fileupload_value(job_upload.value)
            if files:
                job_text = extract_text_from_bytes(files[0]['name'], files[0]['content'])
        if not job_text:
            print(" Please provide a job description (paste or upload).")
            return
        job_text = clean_text(job_text)
        print(" Job description loaded.")

        print(f" Extracting job requirements using {LLM_EXTRACT_MODEL}...")
        requirements = extract_requirement_clusters(job_text, model=LLM_EXTRACT_MODEL)
        
        if not any(requirements.values()):
             print("Extraction failed. Please check OpenRouter key or model availability.")
             return
             
        all_skills = requirements["must_have"] + requirements["important"] + requirements["nice_to_have"]
        print("Extracted requirement clusters:", requirements)

        print(" Detecting working embedding model...")
        embed_model = detect_working_openrouter_embedding_model()
        if not embed_model:
            print(" Could not find a working embedding model. Check your OPENROUTER_API_KEY.")
            return
        print("Using embedding model:", embed_model)
        files = normalize_fileupload_value(resume_upload.value)
        if not files:
            print(" Please upload at least one resume file.")
            return
        candidates = []
        print(f" Processing {len(files)} resumes...")
        for idx, f in enumerate(files):
            print(f"  -> Processing {f['name']}...")
            txt = extract_text_from_bytes(f['name'], f['content'])
            txt = clean_text(txt)
            
            skill_matches = compute_skill_matches_for_resume(txt, all_skills, embed_model)
            

            overall_score = 0.0
            for skill_phrase, v in skill_matches.items():
                sim_score = v['sim']
                weight = 1.0 
                if skill_phrase in requirements['must_have']:
                    weight = 4.0 # High priority
                elif skill_phrase in requirements['important']:
                    weight = 1.5 # Medium priority

                final_score = sim_score * weight
             
                if v['kw_match'] and sim_score >= KEYWORD_BONUS_THRESHOLD:
                     final_score += 0.5 
                
                if sim_score < 0.7:
                    final_score *= 0.1
                    
                overall_score += final_score
          
            matched_skills = [k for k,v in skill_matches.items() if v['kw_match'] or v['sim'] > SEMANTIC_THRESHOLD]
            gaps = [k for k in requirements['must_have'] if k not in matched_skills]
            
            all_evidence = [v['best_evidence'] for v in skill_matches.values() if v['best_evidence'] is not None and v['sim'] > SEMANTIC_THRESHOLD]
            top_evidence = sorted(all_evidence, key=lambda x: x['score'], reverse=True)[:3]

            candidates.append({
                "id": idx+1,
                "name": f['name'],
                "text": txt,
                "overall_score": overall_score,
                "matched_skills": matched_skills,
                "gaps": gaps,
                "top_evidence": top_evidence 
            })

        print("Generating explainable comparisons...")
        explanations = explainable_ranking_pipeline(candidates)

        print("\n" + "="*60)
        print("Generated Explanations (Comparator Reports)")
        print("="*60)
        for exp in explanations:
            print(exp)
            print("-" * 60)

        print("\n Ranking summary:")
        sorted_cands = sorted(candidates, key=lambda x: x['overall_score'], reverse=True)
        for c in sorted_cands:
            print(f"{c['name']} - Score: {c['overall_score']:.3f}, Matched Skills: {', '.join(c['matched_skills'])}")

        
        if REPORTLAB_AVAILABLE:
            try:
                generate_pdf(job_text, requirements, candidates, explanations)
                print("\nPDF report generated! Click the link above to download.")
            except Exception as e:
                print("Failed to generate PDF:", e)
        else:
            print("\nNOTE: Reportlab library is missing. PDF report generation skipped.")

print("Main processing logic defined.")


# UI SETUP

# In[9]:


job_textarea = widgets.Textarea(value="", placeholder="Paste job description (required)", description="Job text:", layout=widgets.Layout(width="100%", height="140px"))
job_upload = widgets.FileUpload(accept=".txt,.pdf,.docx,.csv,.json", multiple=False, description="Upload job file (optional)")
resume_upload = widgets.FileUpload(accept=".txt,.pdf,.docx,.csv,.json", multiple=True, description="Upload resume(s)")
process_btn = widgets.Button(description="Process & Explain", button_style="success")
output = widgets.Output()

print("Widgets created.")


# UI DISPLAY

# In[10]:


process_btn.on_click(on_process_click)
display(widgets.VBox([
    widgets.Label("Paste job description OR upload file. Then upload one or more resumes (txt/pdf/docx/csv/json)."),
    job_textarea, job_upload, widgets.HTML("<hr>"), resume_upload, process_btn, output
]))

print("Interactive UI displayed. Click 'Process & Explain' after uploading files.")


# ENTIRE PROJECT STREAMLIT DEPLOYMENT

# In[2]:


get_ipython().run_cell_magic('writefile', 'app.py', 'import streamlit as st\nimport os, io, re, json, time, hashlib\nfrom collections import defaultdict\nimport numpy as np\nimport pandas as pd\nimport requests\n\ntry:\n    import pdfplumber\nexcept ImportError:\n    pdfplumber = None\ntry:\n    import docx\nexcept ImportError:\n    docx = None\n\ntry:\n    from reportlab.lib.pagesizes import A4\n    from reportlab.pdfgen import canvas\n    from reportlab.lib.units import cm\n    import base64\n    REPORTLAB_AVAILABLE = True\nexcept ImportError:\n    REPORTLAB_AVAILABLE = False\n\n\nOPENROUTER_EMBED_URL = "https://openrouter.ai/api/v1/embeddings"\nOPENROUTER_CHAT_URL = "https://openrouter.ai/api/v1/chat/completions"\nEMBED_CANDIDATES = ["openai/text-embedding-3-large","openai/text-embedding-3-small"]\nLLM_EXTRACT_MODEL = "openai/gpt-3.5-turbo" \nCHAT_MODEL_DEFAULT = "mistralai/mistral-7b-instruct"\n\n\nSEMANTIC_THRESHOLD = 0.75     \nKEYWORD_BONUS_THRESHOLD = 0.7  \n\n@st.cache_data\ndef extract_text_from_uploaded_file(uploaded_file):\n    """Extracts text from various file types supported by Streamlit UploadedFile."""\n    filename = uploaded_file.name\n    file_bytes = uploaded_file.getvalue()\n    \n    fname = filename.lower()\n\n    if fname.endswith(".txt"):\n        return file_bytes.decode("utf-8", errors="replace")\n\n    if fname.endswith(".pdf") and pdfplumber:\n        try:\n            pages=[]\n            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:\n                for p in pdf.pages:\n                    pages.append(p.extract_text() or "")\n            return "\\n".join(pages)\n        except Exception:\n            return ""\n            \n  \n    if fname.endswith(".docx") and docx:\n        try:\n            doc = docx.Document(io.BytesIO(file_bytes))\n            return "\\n".join([p.text for p in doc.paragraphs])\n        except Exception:\n            return ""\n\n    try:\n        return file_bytes.decode("utf-8", errors="replace")\n    except Exception:\n        return ""\n\ndef sent_tokenize(text):\n    text = (text or "").strip()\n    if not text:\n        return []\n    sents = re.split(r\'(?<=[.!?])\\s+\', text)\n    return [s.strip() for s in sents if s.strip()]\n\ndef clean_text(txt):\n    txt = str(txt or "")\n    txt = txt.replace("\\x0c"," ")\n    txt = re.sub(r\'\\(cid:\\d+\\)\',\' \', txt)\n    txt = re.sub(r\'\\s+\',\' \', txt).strip()\n    return txt\n\n\n@st.cache_data(show_spinner="Generating embeddings via OpenRouter...")\ndef get_openrouter_embeddings(texts, model, api_key, batch_size=16):\n    if not api_key:\n        raise RuntimeError("OPENROUTER_API_KEY missing.")\n    \n    headers = {"Authorization": f"Bearer {api_key}", "Content-Type":"application/json"}\n    all_embs = []\n    \n    for i in range(0, len(texts), batch_size):\n        batch = texts[i:i+batch_size]\n        payload = {"model": model, "input": batch}\n        \n        try:\n            resp = requests.post(OPENROUTER_EMBED_URL, headers=headers, json=payload, timeout=60)\n            resp.raise_for_status()\n        except requests.exceptions.RequestException as e:\n            st.error(f"OpenRouter Embeddings Failed: {e}")\n            raise\n\n        j = resp.json()\n        for item in j["data"]:\n            all_embs.append(np.array(item["embedding"], dtype=np.float32))\n        time.sleep(0.05)\n        \n    if not all_embs:\n        return np.zeros((0,0), dtype=np.float32)\n        \n    arr = np.vstack(all_embs)\n    arr = arr / np.linalg.norm(arr, axis=1, keepdims=True)\n    return arr\n\n\n@st.cache_data(show_spinner="Extracting requirements via OpenRouter...")\ndef openrouter_chat(prompt, model, api_key, temperature=0.0, max_tokens=512):\n    if not api_key:\n        raise RuntimeError("OPENROUTER_API_KEY missing.")\n        \n    headers = {"Authorization": f"Bearer {api_key}", "Content-Type":"application/json"}\n    payload = {"model": model, "messages": [{"role":"user","content": prompt}], "temperature": temperature, "max_tokens": max_tokens}\n    \n    try:\n        resp = requests.post(OPENROUTER_CHAT_URL, headers=headers, json=payload, timeout=120)\n        resp.raise_for_status()\n    except requests.exceptions.RequestException as e:\n        st.error(f"OpenRouter Chat Failed: {e}")\n        raise\n        \n    j = resp.json()\n    return j["choices"][0]["message"]["content"].strip()\n\n\n\ndef extract_requirement_clusters(job_text, api_key, model=LLM_EXTRACT_MODEL):\n    prompt = (\n        "You are a concise parser. Given a job description, return a JSON object with keys "\n        "\\"must_have\\", \\"important\\", \\"nice_to_have\\". Each value should be an array of short skill/competency phrases (1-5 words each). "\n        "It is CRITICAL that you extract the actual skills/competencies, not filler words like \'of\' or \'the\'. "\n        "If the job contains explicit labels like \'must have\' or lists, follow them. Keep phrases short. Return only JSON.\\n\\n"\n        f"JOB_DESCRIPTION:\\n{job_text}\\n\\nReturn JSON now."\n    )\n    txt = openrouter_chat(prompt, model=model, api_key=api_key, temperature=0.0, max_tokens=500)\n    m = re.search(r"\\{.*\\}", txt, flags=re.S)\n    \n    if not m:\n        st.warning("LLM failed to return structured JSON. Using token frequency fallback.")\n        words = re.findall(r"\\b[A-Za-z0-9\\-\\+#\\.\\_]+\\b", job_text.lower())\n        common = sorted(set(words), key=lambda w: -words.count(w))[:12]\n        return {"must_have": common[:4], "important": common[4:8], "nice_to_have": common[8:12]}\n        \n    try:\n        obj = json.loads(m.group(0))\n    except json.JSONDecodeError:\n        st.warning("LLM returned invalid JSON. Using token frequency fallback.")\n        words = re.findall(r"\\b[A-Za-z0-9\\-\\+#\\.\\_]+\\b", job_text.lower())\n        common = sorted(set(words), key=lambda w: -words.count(w))[:12]\n        obj = {"must_have": common[:4], "important": common[4:8], "nice_to_have": common[8:12]}\n        \n    for k in ("must_have","important","nice_to_have"):\n        if k not in obj:\n            obj[k] = []\n        if isinstance(obj[k], list):\n            # Clean up short common words\n            obj[k] = [p.strip() for p in obj[k] if len(p.split()) > 1 or p.lower() not in [\'of\', \'and\', \'in\', \'to\', \'the\', \'with\', \'for\', \'or\']]\n            \n    return obj\n\ndef compute_skill_matches_for_resume(resume_text, skill_phrases, embed_model, api_key, chunk_size=3):\n    out = {}\n    r_low = resume_text.lower()\n    \n    for p in skill_phrases:\n        p_low = p.lower()\n        kw_match = bool(re.search(r"\\b" + re.escape(p_low) + r"\\b", r_low))\n        out[p] = {"kw_match": kw_match, "sim": 0.0, "best_evidence": None}\n    \n    sentences = sent_tokenize(resume_text)\n    chunks = [" ".join(sentences[i:i+chunk_size]) for i in range(len(sentences)) if " ".join(sentences[i:i+chunk_size])]\n    \n    if not chunks or not skill_phrases:\n        return out\n        \n    try:\n        texts = skill_phrases + chunks\n        embs = get_openrouter_embeddings(texts, model=embed_model, api_key=api_key, batch_size=8)\n        \n        skill_embs = embs[:len(skill_phrases)]\n        chunk_embs = embs[len(skill_phrases):]\n        \n        sims_matrix = (skill_embs @ chunk_embs.T).astype(float) \n        \n        for i, p in enumerate(skill_phrases):\n            max_sim = np.max(sims_matrix[i])\n            best_chunk_index = np.argmax(sims_matrix[i])\n            \n            out[p]["sim"] = float(max_sim)\n            out[p]["best_evidence"] = {\n                "skill": p,\n                "chunk_text": chunks[best_chunk_index],\n                "score": float(max_sim)\n            }\n            \n    except Exception as e:\n        st.error(f"Error during embedding computation: {e}")\n        pass\n        \n    return out\n\n\ndef generate_candidate_comparator_explanation(candidateA, candidateB):\n    scoreA = candidateA.get(\'overall_score\', 0.0)\n    scoreB = candidateB.get(\'overall_score\', 0.0)\n\n    explanation = (\n        f"**CANDIDATE A** (Name: **{candidateA.get(\'name\')}** | Score: **{scoreA:.3f}**):\\n"\n        f"Matched must-have skills: {\', \'.join(candidateA.get(\'matched_skills\',[]))}\\n"\n        f"Gaps: {\', \'.join(candidateA.get(\'gaps\',[]))}\\n"\n        f"Top evidence snippets:\\n"\n    )\n    for e in candidateA.get(\'top_evidence\', []):\n        explanation += f" - **{e.get(\'skill\', \'Skill\')}:** (Score {e.get(\'score\', 0.0):.3f}) *{e[\'chunk_text\'][:100]}...*\\n"\n\n    explanation += f"\\n**CANDIDATE B** (Name: **{candidateB.get(\'name\')}** | Score: **{scoreB:.3f}**):\\n"\n    explanation += f"Matched must-have skills: {\', \'.join(candidateB.get(\'matched_skills\',[]))}\\n"\n    explanation += f"Gaps: {\', \'.join(candidateB.get(\'gaps\',[]))}\\n"\n    explanation += f"Top evidence snippets:\\n"\n    for e in candidateB.get(\'top_evidence\', []):\n        explanation += f" - **{e.get(\'skill\', \'Skill\')}:** (Score {e.get(\'score\', 0.0):.3f}) *{e[\'chunk_text\'][:100]}...*\\n"\n\n    if scoreA > scoreB:\n        explanation += "\\n**Conclusion: Candidate A is stronger based on overall score.**\\n"\n    elif scoreB > scoreA:\n        explanation += "\\n**Conclusion: Candidate B is stronger based on overall score.**\\n"\n    else:\n        explanation += "\\n**Conclusion: Both candidates have equal scores.**\\n"\n\n    return explanation\n\ndef explainable_ranking_pipeline(candidates):\n    sorted_candidates = sorted(candidates, key=lambda x: x[\'overall_score\'], reverse=True)\n    explanations = []\n    \n    if len(sorted_candidates) >= 2:\n        for i in range(len(sorted_candidates) - 1):\n            a = sorted_candidates[i]\n            b = sorted_candidates[i + 1]\n            explanation = generate_candidate_comparator_explanation(a, b)\n            explanations.append(explanation)\n            \n    return explanations, sorted_candidates\n\n\ndef generate_pdf_report(job_text, requirements, candidates, explanations):\n    if not REPORTLAB_AVAILABLE:\n        st.warning("Reportlab library not found. PDF report generation skipped.")\n        return None\n        \n    buffer = io.BytesIO()\n    c = canvas.Canvas(buffer, pagesize=A4)\n    width, height = A4\n    margin = 2*cm\n    y = height - margin\n\n    def write_line(text, line_height=14, is_bold=False):\n        nonlocal y\n        if y < margin + line_height * 2: # Check if a new page is needed soon\n            c.showPage()\n            y = height - margin\n        \n        if is_bold:\n             c.setFont("Helvetica-Bold", line_height-4)\n        else:\n             c.setFont("Helvetica", line_height-4)\n\n        lines = text.split(\'\\n\')\n        for line in lines:\n            max_width = width - 2 * margin\n            current_line = ""\n            for word in line.split(\' \'):\n                test_line = current_line + word + " "\n                if c.stringWidth(test_line) < max_width:\n                    current_line = test_line\n                else:\n                    c.drawString(margin, y, current_line.strip())\n                    y -= line_height\n                    current_line = word + " "\n            \n            if current_line.strip():\n                c.drawString(margin, y, current_line.strip())\n                y -= line_height\n                \n        c.setFont("Helvetica", 12)\n        return y \n\n\n    c.setFont("Helvetica-Bold", 16)\n    write_line("Candidate Matching Report", line_height=20, is_bold=True)\n    write_line(f"Generated: {time.strftime(\'%Y-%m-%d %H:%M:%S\')}\\n", is_bold=False)\n\n    # Job description\n    write_line("Job Description Summary:", 16, is_bold=True)\n    write_line(job_text[:500] + ("..." if len(job_text) > 500 else ""), 12)\n    \n    write_line("")\n    write_line("Extracted Requirement Clusters:", 16, is_bold=True)\n    for k,v in requirements.items():\n        write_line(f"**{k.replace(\'_\', \' \').title()}:** {\', \'.join(v)}", 12)\n\n    write_line("")\n    write_line("Candidate Comparisons:", 16, is_bold=True)\n    for exp in explanations:\n        write_line(exp.replace(\'**\', \'\').replace(\'\\n\', \'\\n\')) # Remove markdown for PDF text\n        write_line("-" * 50)\n\n    write_line("")\n    write_line("Ranking Summary:", 16, is_bold=True)\n    for cnd in sorted(candidates, key=lambda x: x[\'overall_score\'], reverse=True):\n        write_line(f"RANK {cnd[\'name\']} - Score: {cnd[\'overall_score\']:.3f}")\n        write_line(f"Matched Skills: {\', \'.join(cnd[\'matched_skills\'])}")\n        write_line(f"Gaps: {\', \'.join(cnd[\'gaps\'])}")\n        write_line("-" * 20)\n\n    c.save()\n    buffer.seek(0)\n    return buffer\n\ndef main():\n    st.set_page_config(page_title="Resume Matcher", layout="wide")\n    st.title("📄 Semantic Resume Matcher and Ranker")\n    st.markdown("Use OpenRouter (OpenAI/Mistral) embeddings to compare resumes against a job description. The system heavily weights \'Must-Have\' skills.")\n    st.markdown("---")\n\n   \n    st.sidebar.header("API Configuration")\n    openrouter_api_key = st.sidebar.text_input(\n        "OpenRouter API Key",\n        type="password",\n        value=os.getenv("OPENROUTER_API_KEY", "") # Fallback to env var\n    )\n    if not openrouter_api_key:\n        st.sidebar.error("Please enter your OpenRouter API Key.")\n        return\n\n\n    st.header("1. Job Description (JD)")\n    jd_tab, jd_file_tab = st.tabs(["Paste Text", "Upload File"])\n    \n    with jd_tab:\n        job_text = st.text_area(\n            "Paste Job Description Text here:",\n            height=200,\n            key="jd_text_area"\n        )\n    with jd_file_tab:\n        jd_file = st.file_uploader("Upload Job Description File (.txt, .pdf, .docx)", type=["txt", "pdf", "docx"], key="jd_file_uploader")\n        if jd_file:\n            job_text = extract_text_from_uploaded_file(jd_file)\n\n    if not job_text:\n        st.info("Please input or upload a Job Description to proceed.")\n        return\n\n    st.header("2. Resumes")\n    resume_files = st.file_uploader(\n        "Upload Candidate Resumes (.pdf, .docx, .txt)",\n        type=["pdf", "docx", "txt"],\n        accept_multiple_files=True\n    )\n\n    if not resume_files:\n        st.info("Please upload at least one resume.")\n        return\n\n    st.markdown("---")\n\n    if st.button(" Process Candidates Resume"):\n        \n       \n        st.cache_data.clear()\n        \n        try:\n      \n            cleaned_job_text = clean_text(job_text)\n            st.success("Job Description loaded and cleaned.")\n\n            requirements = extract_requirement_clusters(cleaned_job_text, openrouter_api_key, model=LLM_EXTRACT_MODEL)\n            all_skills = requirements["must_have"] + requirements["important"] + requirements["nice_to_have"]\n            st.subheader("Extracted Requirements:")\n            st.json(requirements)\n\n            candidates = []\n            resume_progress = st.progress(0, text="Processing resumes...")\n            \n            for idx, f in enumerate(resume_files):\n                st.info(f"Processing **{f.name}**...")\n                \n                resume_text = extract_text_from_uploaded_file(f)\n                cleaned_resume_text = clean_text(resume_text)\n\n                skill_matches = compute_skill_matches_for_resume(cleaned_resume_text, all_skills, EMBED_CANDIDATES[0], openrouter_api_key)\n                \n                overall_score = 0.0\n                for skill_phrase, v in skill_matches.items():\n                    sim_score = v[\'sim\']\n                    weight = 1.0 \n                    if skill_phrase in requirements[\'must_have\']:\n                        weight = 4.0 \n                    elif skill_phrase in requirements[\'important\']:\n                        weight = 1.5 \n\n                    final_score = sim_score * weight\n                    \n                    if v[\'kw_match\'] and sim_score >= KEYWORD_BONUS_THRESHOLD:\n                         final_score += 0.5 \n                    \n                    if sim_score < 0.7:\n                        final_score *= 0.1\n                        \n                    overall_score += final_score\n        \n                matched_skills = [k for k,v in skill_matches.items() if v[\'kw_match\'] or v[\'sim\'] > SEMANTIC_THRESHOLD]\n                gaps = [k for k in requirements[\'must_have\'] if k not in matched_skills]\n                \n                all_evidence = [v[\'best_evidence\'] for v in skill_matches.values() if v[\'best_evidence\'] is not None and v[\'sim\'] > SEMANTIC_THRESHOLD]\n                top_evidence = sorted(all_evidence, key=lambda x: x[\'score\'], reverse=True)[:3]\n\n                candidates.append({\n                    "id": idx+1,\n                    "name": f.name,\n                    "overall_score": overall_score,\n                    "matched_skills": matched_skills,\n                    "gaps": gaps,\n                    "top_evidence": top_evidence \n                })\n                resume_progress.progress((idx + 1) / len(resume_files), text=f"Processed **{f.name}**")\n\n            resume_progress.empty()\n\n            explanations, sorted_candidates = explainable_ranking_pipeline(candidates)\n\n            st.header("3. Ranking Summary ")\n        \n            ranking_data = [\n                {"Rank": i+1, "Candidate": c[\'name\'], "Score": f"{c[\'overall_score\']:.3f}", "Matched Skills": ", ".join(c[\'matched_skills\'])}\n                for i, c in enumerate(sorted_candidates)\n            ]\n            st.dataframe(ranking_data, use_container_width=True, hide_index=True)\n            \n            st.header("4. Head-to-Head Comparisons")\n            for exp in explanations:\n                st.markdown(exp)\n                st.markdown("---")\n\n            pdf_buffer = generate_pdf_report(cleaned_job_text, requirements, candidates, explanations)\n            if pdf_buffer:\n                st.download_button(\n                    label="Download PDF Report",\n                    data=pdf_buffer,\n                    file_name="candidate_ranking_report.pdf",\n                    mime="application/pdf"\n                )\n        \n        except RuntimeError as e:\n            st.error(f"A critical error occurred: {e}. Please check your API key and network connection.")\n\nif __name__ == "__main__":\n    main()')


# In[ ]:




